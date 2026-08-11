# -*- coding: utf-8 -*-
"""
 GÉNÉRATEUR DE CV — module de l'agent (compatible Streamlit Cloud)
 Produit un CV Word (.docx) dense et professionnel depuis config/mon_profil.json.
 Peut cibler une offre : compétences citées remontées en tête.
"""
import os, io, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x1F, 0x3B, 0x57)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GREY   = RGBColor(0x5B, 0x70, 0x80)
DARK   = RGBColor(0x20, 0x30, 0x3F)


def _bord_bas(paragraph, color="2E75B6", size="8"):
    p = paragraph._p; pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)


def _run(p, texte, size=10.5, color=DARK, bold=False, italic=False, name="Calibri"):
    r = p.add_run(texte); r.font.size = Pt(size); r.font.color.rgb = color
    r.bold = bold; r.italic = italic; r.font.name = name
    return r


def _titre_section(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
    _run(p, texte.upper(), size=12, color=NAVY, bold=True)
    _bord_bas(p)
    return p


def _competences_ordonnees(profil, offre=None):
    comps = list(profil["competences"])
    if offre:
        txt = (offre.get("description", "") + " " +
               " ".join(offre.get("competences_requises", []))).lower()
        comps.sort(key=lambda c: 0 if c.lower() in txt else 1)
    return comps


def generer_cv(profil, offre=None, dossier=None, en_memoire=False):
    cv = profil.get("cv", {})
    poste = (offre["titre"] if offre else profil["postes_vises"][0])

    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = Inches(0.55); s.bottom_margin = Inches(0.55)
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)

    # EN-TÊTE
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    _run(p, profil["nom"], size=24, color=NAVY, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    _run(p, poste, size=13.5, color=ACCENT, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    contact = f"{profil['localisation']}  ·  {profil['telephone']}  ·  {profil['email']}"
    if profil.get("linkedin"): contact += f"  ·  {profil['linkedin']}"
    _run(p, contact, size=9.5, color=GREY)

    # PROFIL
    _titre_section(doc, "Profil professionnel")
    if cv.get("resume_pro"):
        pr = doc.add_paragraph(); pr.paragraph_format.space_after = Pt(2)
        pr.paragraph_format.line_spacing = 1.05
        _run(pr, cv["resume_pro"], size=10.5, color=DARK)

    # COMPÉTENCES
    _titre_section(doc, "Compétences clés")
    cats = cv.get("competences_categories")
    if cats:
        for c in cats:
            pc = doc.add_paragraph(); pc.paragraph_format.space_after = Pt(1)
            _run(pc, c["categorie"] + " : ", size=10, color=NAVY, bold=True)
            _run(pc, "  •  ".join(c["items"]), size=10, color=DARK)
    else:
        pc = doc.add_paragraph()
        _run(pc, "  •  ".join(_competences_ordonnees(profil, offre)), size=10, color=DARK)

    # EXPÉRIENCE
    _titre_section(doc, "Expérience professionnelle")
    for e in cv.get("experiences", []):
        pe = doc.add_paragraph(); pe.paragraph_format.space_before = Pt(4)
        pe.paragraph_format.space_after = Pt(0)
        _run(pe, e["poste"], size=11, color=NAVY, bold=True)
        _run(pe, "   " + e["entreprise"], size=10.5, color=ACCENT)
        pd = doc.add_paragraph(); pd.paragraph_format.space_after = Pt(1)
        _run(pd, f"{e['periode']}" + (f"  ·  {e['lieu']}" if e.get("lieu") else ""),
             size=9.5, color=GREY, italic=True)
        for t_ in e.get("taches", []):
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.space_after = Pt(1); pb.paragraph_format.line_spacing = 1.02
            _run(pb, t_, size=10, color=DARK)

    # PROJETS
    if cv.get("projets"):
        _titre_section(doc, "Projets Data marquants")
        for pr in cv["projets"]:
            pp = doc.add_paragraph(style="List Bullet"); pp.paragraph_format.space_after = Pt(1)
            _run(pp, pr["titre"] + " — ", size=10, color=NAVY, bold=True)
            _run(pp, pr["detail"], size=10, color=DARK)

    # FORMATION
    _titre_section(doc, "Formation")
    for f in cv.get("formations", []):
        pf = doc.add_paragraph(style="List Bullet"); pf.paragraph_format.space_after = Pt(1)
        _run(pf, f["intitule"], size=10, color=DARK)
        if f.get("annee"): _run(pf, f"   ({f['annee']})", size=9.5, color=GREY)

    # CERTIFICATIONS
    if cv.get("certifications"):
        _titre_section(doc, "Certifications")
        for c in cv["certifications"]:
            pc = doc.add_paragraph(style="List Bullet"); pc.paragraph_format.space_after = Pt(1)
            _run(pc, c["intitule"], size=10, color=DARK)
            if c.get("annee"): _run(pc, f"   ({c['annee']})", size=9.5, color=GREY)

    # ATOUTS & LANGUES
    _titre_section(doc, "Atouts & langues")
    for a in cv.get("atouts", []):
        pa = doc.add_paragraph(style="List Bullet"); pa.paragraph_format.space_after = Pt(1)
        _run(pa, a, size=10, color=DARK)
    pl = doc.add_paragraph(); pl.paragraph_format.space_before = Pt(2)
    _run(pl, "Langues : ", size=10, color=NAVY, bold=True)
    _run(pl, " · ".join(profil.get("langues", [])), size=10, color=DARK)

    if en_memoire:
        buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
    dossier = dossier or os.path.join(os.path.dirname(os.path.abspath(__file__)), "sorties")
    os.makedirs(dossier, exist_ok=True)
    suffixe = ("_" + re.sub(r"[^A-Za-z0-9]", "_", offre["entreprise"])[:15]) if offre and offre.get("entreprise") else ""
    chemin = os.path.join(dossier, f"CV_{re.sub(r'[^A-Za-z0-9]','_',profil['nom'])}{suffixe}.docx")
    doc.save(chemin)
    return chemin


if __name__ == "__main__":
    import json
    prof = json.load(open(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                       "config", "mon_profil.json"), encoding="utf-8"))
    print("CV général :", generer_cv(prof))
